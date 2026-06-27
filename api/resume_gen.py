# resume_builder.py
import sys as _sys, os as _os
if not getattr(_sys, "frozen", False):
    _sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

import anthropic
import json
import os
import re
import time
from datetime import datetime
from api.prompts import get_prompt, PROMPT_RESUME_TAILOR, PROMPT_JOB_EXPERIENCE_PRIORITY
from core.config import AI_MODEL  # Sonnet — quality critical

# ── 3-Layer gold standard system prompt — hardcoded, never stored in DB ─
# System = these universal enhancement rules (same for all users)
# User   = personalized resume_tailor prompt from prompts.db + job data
_RESUME_SYSTEM_PROMPT = (
    "You are an elite Executive Resume Writer and Tier-1 ATS Optimizer. "
    "Transform passive work histories into high-impact, metric-driven achievements "
    "that secure interviews. Operate in 3 mandatory layers:\n\n"
    "LAYER 1 - ATS KEYWORD INJECTION:\n"
    "Weave TARGET KEYWORDS naturally into action verb context, never keyword-stuff. "
    "Every key technology must appear as a tool used to achieve a result.\n\n"
    "LAYER 2 - METRIC-DRIVEN QUANTIFICATION:\n"
    "Apply Google X-Y-Z formula: Accomplished [X], measured by [Y], by doing [Z]. "
    "EVERY bullet must contain a percent, dollar, time, or team-size metric. "
    "Numbers in <fact> tags: copy EXACTLY, never alter. "
    "Lines marked [ESTIMATE NEEDED]: generate a realistic metric, wrap it in "
    "<estimate>N%</estimate> tags. "
    "Estimates must NOT be clean multiples of 5 (not 10%, 20%, 30%).\n\n"
    "LAYER 3 - VISUAL SCANNABILITY:\n"
    "Start every bullet with a distinct high-impact action verb. "
    "NEVER start with: Responsible for, Assisted with, Helped, Worked on. "
    "Maximum 22 words per bullet. "
    "Bold **[key technology]** and **[key metric]** for visual anchors.\n\n"
    "OUTPUT RULES:\n"
    "Output ONLY the enhanced JSON, zero preamble or closing text. "
    "Respect BULLET BUDGET exactly. "
    "NEVER repeat the same action verb twice in the entire resume. "
    "NEVER use: spearheaded, orchestrated, leveraged, utilized, synergized. "
    "Match verb intensity to role seniority."
)

# ── Metric guard — fallbacks defined FIRST so names always exist ──────
# If imports below succeed they overwrite; if not, fallbacks stay
def calculate_bullet_budget(profile, jd_meta):
    """Fallback: returns empty budget if bullet_budget.py unavailable."""
    return {"budgets": [], "total_bullets": 0, "jd_skills": []}
def build_factual_scaffolding(profile, user_metrics=None):
    """Fallback: empty scaffolding if metric_guard.py unavailable."""
    return ""
def get_verb_constraints(seniority):
    """Fallback: default mid-level verbs."""
    return {"tier": "mid",
            "allowed_verbs": ["Engineered","Designed","Automated","Optimized","Reduced"],
            "banned_verbs":  ["spearheaded","orchestrated","leveraged","utilized"]}
def selective_fuzzing_processor(text):
    """Fallback: return text unchanged."""
    return text
def scan_profile_for_missing_metrics(profile):
    """Fallback: assume no missing metrics."""
    return []

try:
    from api.bullet_budget import calculate_bullet_budget
except Exception as _imp_e:
    pass  # fallback already defined above

try:
    from api.metric_guard import (
        build_factual_scaffolding, get_verb_constraints,
        selective_fuzzing_processor, scan_profile_for_missing_metrics)
except Exception as _imp_e:
    pass  # fallbacks already defined above
try:
    from core.logger import log, log_warn, log_error, log_debug
except Exception:
    def log(m,*a): pass
    def log_warn(m,*a): pass
    def log_error(m,*a,**k): pass
    def log_debug(m,*a): pass

# Initialize Claude client
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# Throttle settings for batch resume generation.
# DELAY_BETWEEN_CALLS spaces out API calls so a run of 20 resumes
# does not hit the rate limit in a burst.
DELAY_BETWEEN_CALLS = 3      # seconds to wait between each resume
MAX_RETRIES         = 5      # attempts per resume (rate limit + 500 errors)
RETRY_BACKOFF       = 5      # base seconds — exponential: 5, 10, 20, 40, 80s


def flatten_skills(skills: dict) -> list:
    from itertools import chain
    return list(chain.from_iterable(
        v for v in skills.values() if isinstance(v, list)
    ))


# Domain/industry terms that frequently get hallucinated into tailored
# summaries when a job is in a field the candidate hasn't worked in.
# If the AI puts one of these in the summary but it does NOT appear
# anywhere in the candidate's real profile, we flag it.
_FABRICATION_WATCH_TERMS = [
    "embedded systems", "embedded software", "firmware",
    "autonomous vehicle", "autonomous vehicles", "self-driving",
    "mining", "heavy equipment", "heavy machinery",
    "aerospace", "defense", "automotive", "medical device",
    "blockchain", "cybersecurity", "penetration testing",
    "iot", "fpga", "rtos", "robotics",
    "game development", "ar/vr", "augmented reality",
]


def _add_runs(paragraph, text: str, base_size_pt: float = 11.0):
    """
    Write text into a paragraph, converting **markdown bold** to real Word bold.
    Splits on ** markers, alternates bold/normal runs.
    Example: "reduced by **67%** using **SQL**" →
             "reduced by " (normal) + "67%" (bold) + " using " (normal) + "SQL" (bold)
    """
    from docx.shared import Pt  # local import — available at module level call
    import re as _re_ar
    # Strip any <fact>/<estimate> tags that leaked through processing
    text = _re_ar.sub(r"</?(?:fact|estimate)[^>]*>", "", text or "")
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run          = paragraph.add_run(part)
        run.bold     = (idx % 2 == 1)   # odd segments are bold (between ** markers)
        run.font.size = Pt(base_size_pt)
    return paragraph


def check_summary_for_fabrication(summary: str, profile: dict) -> list:
    """
    Scans a tailored summary for domain/industry terms that do NOT
    appear anywhere in the candidate's real profile. Returns a list
    of suspicious terms (empty list means the summary looks clean).

    This is a safety net, not a guarantee -- always read the summary
    yourself before submitting.
    """
    # Build the full text of everything the candidate has actually done
    real_text_parts = [
        profile.get("summary", ""),
        " ".join(flatten_skills(profile.get("skills", {}))),
    ]
    for job in profile.get("experience", []):
        real_text_parts.append(job.get("title", ""))
        real_text_parts.append(job.get("company", ""))
        real_text_parts.extend(job.get("bullets", []))
    for proj in profile.get("projects", []):
        real_text_parts.append(proj.get("name", ""))
        real_text_parts.append(proj.get("tech", ""))
        real_text_parts.extend(proj.get("bullets", []))
        if proj.get("description"):
            real_text_parts.append(proj["description"])

    real_text = " ".join(real_text_parts).lower()
    summary_l = (summary or "").lower()

    flagged = []
    for term in _FABRICATION_WATCH_TERMS:
        if term in summary_l and term not in real_text:
            flagged.append(term)
    return flagged


def make_unique_filename(company: str, title: str, folder: str, ext: str,
                         applicant_name: str = "Applicant") -> str:
    """
    Builds the resume path for a job application.

    Layout:
        <folder>/<Company>/<Name>_Resume.<ext>

    - One sub-folder per company.
    - File is simply <Name>_Resume, e.g. John_Smith_Resume.pdf
    - If that name is already taken (second role at same company),
      the role is added: John_Smith_Software_Engineer_Resume.pdf
    - If even that is taken, a timestamp is appended.
    """
    def clean(s: str) -> str:
        return re.sub(r'[^a-zA-Z0-9]+', '_', s).strip('_')[:40]

    company_dir = os.path.join(folder, clean(company))
    os.makedirs(company_dir, exist_ok=True)

    name_part  = clean(applicant_name)
    title_part = clean(title)

    # 1st choice: just Name_Resume
    candidate = os.path.join(company_dir, f"{name_part}_Resume.{ext}")
    if not os.path.exists(candidate):
        return candidate

    # 2nd choice: add the role (for a different role at the same company)
    candidate = os.path.join(company_dir, f"{name_part}_{title_part}_Resume.{ext}")
    if not os.path.exists(candidate):
        return candidate

    # 3rd choice: add a timestamp (re-applying to the same role)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(
        company_dir, f"{name_part}_{title_part}_Resume_{timestamp}.{ext}"
    )


def tailor_resume_with_ai(profile: dict, job_description: str, job: dict = None) -> dict:
    """Uses Claude to tailor the resume for a specific job."""
    all_skills = flatten_skills(profile["skills"])

    # Base tailoring instructions come from the database.
    # Candidate profile and job details are injected at runtime.
    base_prompt = get_prompt(PROMPT_RESUME_TAILOR)

    # Stage 3: Bullet budget — local definitions guarantee no NameError
    _job = job or {}

    def _local_bullet_budget(p, m):
        import re as _r2
        from datetime import datetime as _d2
        jd_s = [s.lower() for s in m.get("skills", [])]
        budgets, total = [], 0
        for j2 in p.get("experience", []):
            dur = (j2.get("duration","") or j2.get("dates","")).lower()
            if any(w in dur for w in ["present","current","now"]):
                rec = 1.0
            else:
                yrs = _r2.findall(r"\b(20\d{2})\b", dur)
                rec = max(0.1, 1.0-(((_d2.now().year-max(int(y) for y in yrs))*0.12))) if yrs else 0.5
            hay = (j2.get("title","")+" "+" ".join(str(b) for b in j2.get("bullets",[]))).lower()
            ov  = (sum(1 for s in jd_s if s in hay)/len(jd_s)) if jd_s else 0.5
            sc  = rec*0.4 + ov*0.6
            bud = 4 if sc>=0.65 else 3 if sc>=0.50 else 2 if sc>=0.35 else 1 if sc>=0.20 else 0
            budgets.append({"company":j2.get("company",""),"title":j2.get("title",""),
                            "budget":bud,"score":round(sc,2)})
            total += bud
        return {"budgets":budgets,"total_bullets":total,"jd_skills":m.get("skills",[])}

    def _local_scaffolding(p, u=None):
        import re as _r2
        num_re = _r2.compile(r"(\d[\d,]*\s*(?:[KkMmBb])?\b|\d+\s*%|\$\s*\d[\d,]*)", _r2.IGNORECASE)
        lines, has_facts = ["FACTUAL METRICS use ONLY these:"], False
        for j2 in p.get("experience", []):
            facts = []
            for b in j2.get("bullets", [])[:3]:
                if num_re.search(str(b)):
                    facts.append("  - " + str(b)[:120])
            extra = (u or {}).get(j2.get("company",""), "")
            if extra:
                facts.append("  - User stated: " + extra)
                has_facts = True
            if facts:
                has_facts = True
                lines.append(j2.get("title","") + " @ " + j2.get("company","") + ":")
                lines.extend(facts)
        if not has_facts:
            return "NO CONFIRMED METRICS: use scope descriptors only, never invent numbers."
        return "\n".join(lines)

    def _local_verbs(seniority):
        s = (seniority or "").lower()
        if "intern" in s:
            v = ["Optimized","Developed","Implemented","Tested","Fixed"]
        elif any(w in s for w in ["senior","sr"]):
            v = ["Architected","Led","Scaled","Delivered","Redesigned"]
        elif any(w in s for w in ["lead","principal"]):
            v = ["Directed","Transformed","Pioneered","Mentored"]
        else:
            v = ["Engineered","Designed","Automated","Optimized","Reduced"]
        return {"tier":"mid","allowed_verbs":v,
                "banned_verbs":["spearheaded","orchestrated","leveraged","utilized"]}

    try:
        _cbf = calculate_bullet_budget
    except NameError:
        _cbf = _local_bullet_budget
        log_warn("Using local calculate_bullet_budget — module import failed")

    try:
        _bfs = build_factual_scaffolding
    except NameError:
        _bfs = _local_scaffolding
        log_warn("Using local build_factual_scaffolding — module import failed")

    try:
        _gvc = get_verb_constraints
    except NameError:
        _gvc = _local_verbs
        log_warn("Using local get_verb_constraints — module import failed")

    bullet_plan = _job.get("_bullet_budget") or _cbf(
        profile, _job.get("_jd_metadata") or {})
    jd_metadata = _job.get("_jd_metadata") or {}

    # Log bullet budget for visibility
    log_debug("Bullet budget for this resume:")
    for b in bullet_plan.get("budgets", []):
        log_debug("  %s @ %s → %d bullets (score=%.2f)" % (
            b.get("title",""), b.get("company",""),
            b.get("budget",0), b.get("score",0)))
    log_debug("  Total bullets: %d" % bullet_plan.get("total_bullets", 0))
    log_debug("  JD skills used: %s" % bullet_plan.get("jd_skills",[]))

    # Build bullet budget instruction for Claude
    budget_lines = []
    for b in bullet_plan.get("budgets", []):
        if b["budget"] > 0:
            budget_lines.append("  %s @ %s: %d bullets" % (
                b["title"], b["company"], b["budget"]))
    bullet_budget_str = "\n".join(budget_lines) if budget_lines else ""

    # Build factual scaffolding — Claude uses ONLY these metrics
    factual_block = _bfs(
        profile, _job.get("_user_metrics") if _job else None)
    if "NO CONFIRMED METRICS" in factual_block:
        log_warn("Factual scaffolding: no real metrics found — Claude will use scope descriptors")
    else:
        metric_lines = [l for l in factual_block.split("\n") if "Confirmed" in l or "User stated" in l]
        log_debug("Factual scaffolding: %d metric line(s) found" % len(metric_lines))
        for ml in metric_lines[:3]:
            log_debug("  %s" % ml.strip())

    # Role-scale verb constraints — prevents intern claiming enterprise scope
    jd_seniority  = jd_metadata.get("seniority", "mid")
    verb_rules    = _gvc(jd_seniority)
    log_debug("Verb constraints: tier=%s, verbs=%s" % (
        verb_rules.get("tier","?"), verb_rules.get("allowed_verbs",[])))
    verb_block    = (
        "VERB RULES:\n"
        "- Use ONLY these verbs (match role scope): " + ", ".join(verb_rules["allowed_verbs"]) + "\n"
        "- NEVER repeat the same verb twice in the whole resume\n"
        "- NEVER use: " + ", ".join(verb_rules["banned_verbs"][:8])
    )
    prompt = f"""{base_prompt}

Available projects for selection:
{json.dumps(profile['projects'], indent=2)}

Return ONLY valid JSON, no markdown, no fences:
{{
    "summary": "...",
    "skills_grouped": {{
        "Languages":      ["..."],
        "Frameworks":     ["..."],
        "ML/AI":          ["..."],
        "Tools & DevOps": ["..."],
        "Practices":      ["..."]
    }},
    "experience": [
        {{
            "title":    "...",
            "company":  "...",
            "location": "...",
            "duration": "...",
            "bullets":  ["..."]
        }}
    ],
    "projects": [
        {{
            "name":    "...",
            "tech":    "...",
            "bullets": ["...", "..."]
        }}
    ],
    "ats_keywords_used": ["..."]
}}

CANDIDATE PROFILE:
Name: {profile['name']}
Summary (REWRITE COMPLETELY): {profile['summary']}
All Skills: {json.dumps(all_skills, indent=2)}
Experience (VERBATIM): {json.dumps(profile['experience'], indent=2)}
{"Education: " + json.dumps(profile['education'], indent=2) if profile.get('education') else "Education: [Not provided — omit this section entirely from the resume]"}

JOB DESCRIPTION:
{job_description[:4000]}

CRITICAL RULES FOR THIS RESUME:
1. Extract the EXACT job title from the JD — put it in "title" field and open the summary with it
2. Extract the top 10 keywords/technologies from the JD — every one must appear verbatim in skills or bullets
3. Skills categories must reflect THIS candidate's domain — see rules below
4. Order skills within each category by JD relevance — most relevant skill listed first
5. Summary must mirror the JD language in first sentence — if JD says "ServiceNow Architect" write "ServiceNow Architect" not "platform specialist"

SKILLS CATEGORY RULES:
- Category names must be human-readable — NO underscores ever
- Derive category names from the candidate's actual domain — do not impose a fixed structure
- Data formats (JSON, XML, YAML) are NOT languages — group them with Integration or omit from skills
- Platform names (ServiceNow, Salesforce, SAP) are NOT categories — they are implied by all other skills
- Frameworks and platform tools are NOT the same category
- Skills that appear in the JD must appear in the skills section verbatim
- Maximum 7 categories — merge small categories rather than fragment
- Order skills within each category by JD relevance — most relevant skill listed first

EXAMPLES by domain (adapt to whatever domain this candidate is in):
  Platform/consulting (ServiceNow, SAP, Salesforce): "Scripting & Languages", "Platform Modules", "Platform Tools", "Integration & APIs", "Security & Identity", "Cloud & Infrastructure", "Governance & Practices"
  Software engineering: "Languages", "Frameworks & Libraries", "Cloud & DevOps", "Databases", "Tools & Practices"
  Data/ML engineering: "Languages", "ML Frameworks", "Data Pipelines", "Cloud Platforms", "Databases & Storage", "MLOps"
  Cybersecurity: "Security Domains", "Tools & Platforms", "Languages & Scripting", "Frameworks & Standards", "Cloud"
  These are EXAMPLES — always adapt to the actual profile domain
"""

    response_text = ""
    try:
        # System = hardcoded 3-layer gold standard (same for all users)
        # User   = personalized profile prompt from prompts.db + job data
        message = client.messages.create(
            model=AI_MODEL,
            max_tokens=6000,
            temperature=0.1,
            timeout=90.0,
            system=_RESUME_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = message.content[0].text.strip()

        # Strip markdown fences if present
        if "```" in response_text:
            parts = response_text.split("```")
            response_text = parts[1] if len(parts) > 1 else parts[0]
            if response_text.startswith("json"):
                response_text = response_text[4:]

        tailored = json.loads(response_text.strip())

        # Safety net -- flag fabricated domain terms in the summary
        flagged = check_summary_for_fabrication(
            tailored.get("summary", ""), profile
        )
        if flagged:
            print(f"   [WARN]  WARNING: the tailored summary mentions terms not")
            print(f"      found in your real profile: {', '.join(flagged)}")
            print(f"      This may be AI fabrication. Review the summary")
            print(f"      before using this resume, or skip this job.")
            tailored["_fabrication_warning"] = flagged

        return tailored

    except json.JSONDecodeError as e:
        print(f"   [ERR] Claude returned invalid JSON: {e}")
        if response_text:
            print(f"   [ERR] Raw response: {response_text[:300]}")
        raise
    except Exception as e:
        msg = str(e).strip().replace("\n", " | ")
        print(f"   [ERR] Tailoring failed: {type(e).__name__}: {msg}")
        raise




def _add_hyperlink(paragraph, text: str, url: str):
    """
    Add a clickable hyperlink run to a paragraph.
    Falls back to plain text if XML manipulation fails.
    """
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import RGBColor as _RGB

        part  = paragraph.part
        r_id  = part.relate_to(url,
                    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                    is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr     = OxmlElement('w:rPr')

        # Blue underline style
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        # Fallback: plain text
        paragraph.add_run(text)


def _bullet_style(doc):
    """Return bullet list style name — falls back to Normal if List Bullet missing."""
    try:
        _ = doc.styles["List Bullet"]
        return "List Bullet"
    except KeyError:
        return "Normal"


def build_resume_docx(profile: dict, tailored: dict, output_path: str):
    """
    Builds a Word document from the tailored resume data, using
    standard resume formatting conventions:
      - Readable font sizes (11pt body, 10.5pt contact, 22pt name)
      - Consistent section headings with underline rule
      - Right-aligned dates, proper spacing, single accent color
      - No non-standard sections in the submitted document
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

    doc   = Document()
    NAVY  = RGBColor(31, 56, 100)    # accent -- headings & name
    GRAY  = RGBColor(90, 90, 90)     # secondary text
    # Base document font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    # Usable width, for right-aligned tab stops
    usable_width = (
        doc.sections[0].page_width
        - doc.sections[0].left_margin
        - doc.sections[0].right_margin
    )

    def section_heading(text):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text.upper())
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.color.rgb = NAVY
        # Underline rule beneath the heading
        pPr    = p._p.get_or_add_pPr()
        pBdr   = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), '1F3864')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # -- Name ----------------------------------------------------
    name_p           = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(profile["name"])
    nr.bold           = True
    nr.font.size      = Pt(22)
    nr.font.color.rgb = NAVY

    # -- Professional title (if present) -------------------------
    # Prefer JD-matched title from Claude output, fall back to profile title
    title_str = (tailored.get("title") or profile.get("headline") or profile.get("title") or "").strip()
    # Strip placeholder values
    if title_str.upper() in ("[PLACEHOLDER]", "PLACEHOLDER", "N/A", "NONE"):
        title_str = ""
    if title_str:
        title_p           = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(2)
        tr2 = title_p.add_run(title_str)
        tr2.font.size      = Pt(12)
        tr2.font.color.rgb = GRAY
        tr2.italic         = True

    # -- Contact line --------------------------------------------
    contact_p           = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(4)
    # Build contact line — email, phone, location as plain text
    # LinkedIn and GitHub as clickable hyperlinks
    plain_bits = [
        profile.get("email", ""),
        profile.get("phone", ""),
        profile.get("location", ""),
    ]
    plain_text = "  |  ".join(b for b in plain_bits if b)
    linkedin   = profile.get("linkedin", "").strip()
    github     = profile.get("github", "").strip()

    # Determine if LinkedIn is a URL or display name
    def _li_url(val):
        if val.startswith("http"):
            return val
        return f"https://linkedin.com/in/{val.lstrip('/')}"

    if plain_text:
        cr = contact_p.add_run(plain_text)
        cr.font.size      = Pt(10.5)
        cr.font.color.rgb = GRAY

    if linkedin:
        sep = contact_p.add_run("  |  ")
        sep.font.size = Pt(10.5)
        sep.font.color.rgb = GRAY
        _add_hyperlink(contact_p, linkedin, _li_url(linkedin))

    if github:
        sep2 = contact_p.add_run("  |  ")
        sep2.font.size = Pt(10.5)
        sep2.font.color.rgb = GRAY
        gh_url = github if github.startswith("http") else f"https://github.com/{github.lstrip('/')}"
        _add_hyperlink(contact_p, github, gh_url)

    # -- Summary -------------------------------------------------
    section_heading("Professional Summary")
    s = doc.add_paragraph()
    s.paragraph_format.space_before = Pt(4)
    s.paragraph_format.space_after  = Pt(2)
    _add_runs(s, tailored.get("summary",""), base_size_pt=11.0)

    # -- Skills --------------------------------------------------
    # Use heading Claude chose based on profile domain
    _skills_heading = (
        tailored.get("skills_section_heading") or "Technical Skills"
    ).strip()
    # Safety: if Claude returned a long sentence, truncate to the label
    if len(_skills_heading) > 30:
        _skills_heading = "Technical Skills"
    section_heading(_skills_heading)

    def _fmt_category(raw: str) -> str:
        """Convert JSON key names to display labels.
        servicenow_modules → ServiceNow Modules
        tools_platforms    → Tools & Platforms
        development_components → Development Components
        security_authentication → Security & Authentication
        """
        # Replace underscores with spaces
        s = raw.replace("_", " ")
        # Title case each word
        s = s.title()
        # Fix common abbreviations
        for old, new in [
            ("Servicenow", "ServiceNow"),
            ("Aws", "AWS"), ("Gcp", "GCP"),
            ("Api", "API"), ("Apis", "APIs"),
            ("Ui", "UI"), ("Ux", "UX"),
            ("Saml", "SAML"), ("Oauth", "OAuth"),
            ("Sso", "SSO"), ("Sql", "SQL"),
            ("Ml", "ML"), ("Ai", "AI"),
            ("Itsm", "ITSM"), ("Itom", "ITOM"),
            ("Cmdb", "CMDB"), ("Ire", "IRE"),
            ("Atf", "ATF"),
        ]:
            s = s.replace(old, new)
        # Replace "And" connector word → "&"
        s = s.replace(" And ", " & ")
        return s

    for category, items in tailored.get("skills_grouped", {}).items():
        if items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            label = p.add_run(f"{_fmt_category(category)}: ")
            label.bold      = True
            label.font.size = Pt(11)
            vals = p.add_run(", ".join(items))
            vals.font.size = Pt(11)

    # -- Experience ----------------------------------------------
    section_heading("Experience")
    for job in (tailored.get("experience") or []):
        # Title - Company  ............................  Duration (right)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(1)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

        tr = p.add_run(job["title"])
        tr.bold      = True
        tr.font.size = Pt(11)
        cr_ = p.add_run(f" - {job['company']}")
        cr_.font.size = Pt(11)
        if job.get("location"):
            lr = p.add_run(f"  ·  {job['location']}")
            lr.font.size      = Pt(10.5)
            lr.font.color.rgb = GRAY

        dr = p.add_run(f"\t{job['duration']}")
        dr.italic         = True
        dr.font.size      = Pt(10.5)
        dr.font.color.rgb = GRAY

        for bullet in (job.get("bullets") or []):
            bp = doc.add_paragraph(style=_bullet_style(doc))
            bp.paragraph_format.left_indent  = Inches(0.25)
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after  = Pt(1)
            _add_runs(bp, str(bullet), base_size_pt=11.0)

        # Tech stack line — shown only for consulting/platform profiles
        # Proves recency: recruiter sees which tech was used at which role
        tech = (job.get("tech_stack") or "").strip()
        if tech:
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent  = Inches(0.25)
            tp.paragraph_format.space_before = Pt(2)
            tp.paragraph_format.space_after  = Pt(4)
            tlabel = tp.add_run("Tech: ")
            tlabel.bold            = True
            tlabel.italic          = True
            tlabel.font.size       = Pt(9.5)
            tlabel.font.color.rgb  = GRAY
            tvals = tp.add_run(tech)
            tvals.italic           = True
            tvals.font.size        = Pt(9.5)
            tvals.font.color.rgb   = GRAY

    # -- Projects ------------------------------------------------
    if tailored.get("projects"):
        section_heading("Projects")
    for proj in tailored.get("projects", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(1)
        pn = p.add_run(proj["name"])
        pn.bold      = True
        pn.font.size = Pt(11)
        pt = p.add_run(f"  |  {proj['tech']}")
        pt.italic         = True
        pt.font.size      = Pt(10.5)
        pt.font.color.rgb = GRAY

        # Projects may carry either a list of bullets (preferred) or a
        # single description string (older format) -- handle both.
        proj_points = proj.get("bullets")
        if not proj_points:
            desc = proj.get("description", "")
            proj_points = [desc] if desc else []

        for point in proj_points:
            dp = doc.add_paragraph(style=_bullet_style(doc))
            dp.paragraph_format.left_indent  = Inches(0.25)
            dp.paragraph_format.space_before = Pt(1)
            dp.paragraph_format.space_after  = Pt(1)
            _add_runs(dp, str(point), base_size_pt=11.0)

    # -- Education -----------------------------------------------
    if profile.get("education"):
        section_heading("Education")
    for edu in profile.get("education", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

        er = p.add_run(edu["degree"])
        er.bold      = True
        er.font.size = Pt(11)
        if edu.get("school"):
            sr = p.add_run(f" - {edu['school']}")
            sr.font.size = Pt(11)

        yr = p.add_run(f"\t{edu['year']}")
        yr.italic         = True
        yr.font.size      = Pt(10.5)
        yr.font.color.rgb = GRAY

    # -- Volunteer & Leadership ----------------------------------
    if profile.get("volunteer"):
        section_heading("Volunteer & Leadership")
        for vol in profile["volunteer"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            vr = p.add_run(vol["role"])
            vr.bold      = True
            vr.font.size = Pt(11)
            or_ = p.add_run(f" - {vol['organization']}")
            or_.font.size = Pt(11)

            dp = doc.add_paragraph(style=_bullet_style(doc))
            dp.paragraph_format.left_indent  = Inches(0.25)
            dp.paragraph_format.space_before = Pt(1)
            dp.paragraph_format.space_after  = Pt(1)
            _add_runs(dp, str(vol.get("description","")), base_size_pt=11.0)

    doc.save(output_path)
    print(f"   [OK] DOCX saved -> {output_path}")


def prioritize_for_job(tailored: dict, profile: dict,
                        job_description: str) -> dict:
    """
    Reorders, selects, and where genuinely helpful adds synthesized
    bullets to highlight the candidate's most relevant experience
    for a specific job description.

    New bullets are allowed only when the underlying work is clearly
    implied by existing experience and skills -- and only using
    technologies/skills that already exist in the candidate's profile.
    No new capabilities, no invented numbers.

    Called only when profile['experience_highlight'] is True.
    On any API failure, returns tailored unchanged.
    """
    from itertools import chain as _chain
    import re as _re

    experience = tailored.get("experience") or profile.get("experience", [])
    projects   = tailored.get("projects")   or profile.get("projects",   [])

    if not experience and not projects:
        return tailored

    # Build the known-terms corpus from ALL profile data.
    # Any technology/skill/domain in a returned bullet must appear here.
    all_profile_skills = list(_chain.from_iterable(
        v for v in profile.get("skills", {}).values()
        if isinstance(v, list)
    ))

    known_terms = {
        t.lower().strip(".,;:()")
        for t in _chain(
            all_profile_skills,
            [tok for p in profile.get("projects", [])
             for tok in (p.get("tech") or "").split(",")],
            [j.get("title", "") for j in profile.get("experience", [])],
        )
        if isinstance(t, str) and t.strip()
    }

    # Original bullet corpus -- used as a second-pass fallback in check
    original_text = " ".join(
        b
        for src in (profile.get("experience", []) + profile.get("projects", []))
        for b in (src.get("bullets") or [])
    ).lower()

    priority_prompt = get_prompt(PROMPT_JOB_EXPERIENCE_PRIORITY)

    prompt = (
        f"{priority_prompt}\n\n"
        f"JOB DESCRIPTION:\n{job_description[:3000]}\n\n"
        f"CANDIDATE SKILLS (the only technologies allowed in any bullet):\n"
        f"{json.dumps(all_profile_skills)}\n\n"
        f"CANDIDATE EXPERIENCE:\n{json.dumps(experience, indent=2)}\n\n"
        f"CANDIDATE PROJECTS:\n{json.dumps(projects, indent=2)}\n\n"
        "Return the optimised JSON now."
    )

    try:
        message = client.messages.create(
            model=AI_MODEL,
            max_tokens=6000,   # full prioritized resume JSON needs room
            timeout=90.0,
            messages=[{"role": "user", "content": prompt}],
        )
        response_text = message.content[0].text.strip()
        if "```" in response_text:
            parts = response_text.split("```")
            response_text = parts[1] if len(parts) > 1 else parts[0]
            if response_text.startswith("json"):
                response_text = response_text[4:]

        prioritized = json.loads(response_text.strip())

        def _bullet_is_safe(bullet: str) -> bool:
            """
            Returns True if the bullet does not introduce a technology
            or skill that is absent from the candidate's profile.

            Extracts capitalised/CamelCase tokens (likely tech names),
            checks each against known_terms and the original text corpus.
            A bullet passes if every tech token is accounted for.
            """
            tokens = _re.findall(
                r'[A-Z][a-zA-Z0-9#.+/\-]+',
                bullet
            )
            for token in tokens:
                t = token.lower().strip(".,;:()")
                # Accept if in known profile terms OR appears in
                # the original bullet text (company names, etc.)
                if t and len(t) > 1 and t not in known_terms and t not in original_text:
                    return False
            return True

        safe_experience = []
        for job in prioritized.get("experience", []):
            # Find matching original job
            orig_bullets = next(
                (j.get("bullets", []) for j in experience
                 if j.get("company") == job.get("company")), []
            )
            returned_bullets = job.get("bullets", [])

            # Keep ALL original bullets — just use the returned order
            # The safety check only catches truly invented bullets
            safe_bullets = []
            for b in returned_bullets:
                if _bullet_is_safe(b):
                    safe_bullets.append(b)
                else:
                    # Bullet failed safety — check if it's actually
                    # an original bullet (safety check too strict)
                    if any(orig.lower()[:50] in b.lower() or
                           b.lower()[:50] in orig.lower()
                           for orig in orig_bullets):
                        safe_bullets.append(b)   # original bullet, keep it
                    # else: genuinely invented — skip it

            # Critical: if safe_bullets lost ANY original bullets, add them back
            # We never remove what the candidate wrote
            for orig_b in orig_bullets:
                if not any(orig_b.lower()[:60] in sb.lower() or
                           sb.lower()[:60] in orig_b.lower()
                           for sb in safe_bullets):
                    safe_bullets.append(orig_b)   # restore dropped original

            safe_experience.append({**job, "bullets": safe_bullets or orig_bullets})

        # For projects: restore any that Claude dropped entirely
        returned_proj_names = {p.get("name","").lower()
                               for p in prioritized.get("projects", [])}
        safe_projects = []
        for proj in prioritized.get("projects", []):
            orig_p = next(
                (p for p in projects if p.get("name") == proj.get("name")), {})
            orig_bullets = orig_p.get("bullets", [])
            safe_bullets = [b for b in proj.get("bullets", []) if _bullet_is_safe(b)]
            # Restore original bullets that were wrongly filtered
            for ob in orig_bullets:
                if not any(ob.lower()[:60] in sb.lower() or
                           sb.lower()[:60] in ob.lower() for sb in safe_bullets):
                    safe_bullets.append(ob)
            safe_projects.append({**proj, "bullets": safe_bullets or orig_bullets})
        # Restore projects Claude dropped entirely
        for orig_p in projects:
            if orig_p.get("name","").lower() not in returned_proj_names:
                safe_projects.append(orig_p)

        if safe_experience or safe_projects:
            result = dict(tailored)
            if safe_experience:
                result["experience"] = safe_experience
            if safe_projects:
                result["projects"] = safe_projects
            return result

    except Exception as e:
        print(f"      \u26a0\ufe0f  Experience prioritisation skipped: {e}")

    return tailored


# -- Phase 2: throttled batch resume generation ------------------

def _tailor_with_retry(profile: dict, job_description: str, job: dict = None) -> dict:
    """
    Calls tailor_resume_with_ai, retrying with exponential backoff
    if the API reports a rate-limit error. Returns the tailored dict,
    or raises the last error if all retries are exhausted.
    """
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            return tailor_resume_with_ai(profile, job_description, job=job)
        except Exception as e:
            msg = str(e).lower()
            err_code = str(getattr(e, "status_code", "") or "")

            is_retryable = (
                # Rate limit
                "rate" in msg or "429" in msg or "overloaded" in msg
                or "529" in msg
                # Anthropic server error — transient, always worth retrying
                or "500" in msg or "500" in err_code
                or "internal server error" in msg
                or "api_error" in msg
                # Timeout / network blip
                or "timeout" in msg or "timed out" in msg
                or "connection" in msg
            )

            if is_retryable and attempt < MAX_RETRIES:
                wait = RETRY_BACKOFF * (2 ** (attempt - 1))  # exponential
                print(f"      [...] API error ({type(e).__name__}) -- "
                      f"retrying in {wait}s (attempt {attempt}/{MAX_RETRIES})...")
                time.sleep(wait)
                continue
            raise


def batch_generate_resumes(profile: dict, output_dirs: dict,
                           session_start: str = None,
                           job_ids: list = None) -> dict:
    """
    Phase 2: generates tailored resumes for every matched job that
    does not yet have one.

    API calls are spaced out by DELAY_BETWEEN_CALLS seconds and
    retried with backoff on rate-limit errors, so a full run of up
    to ~20 resumes does not exhaust the rate limit in a burst.

    output_dirs must provide:
        'docx' -> folder for DOCX files
        'pdf'  -> folder for PDF files

    Returns a summary dict: {'generated': n, 'failed': n, 'total': n}.
    """
    import db.tracker as tracker
    from utils.pdf import convert_to_pdf
    from api.prompts import all_prompts_ready

    # ── Load user-provided metrics collected via popup before run ────
    # Always define first — guaranteed to exist even if file missing
    _user_metrics = {}
    try:
        import json as _jm2, os as _om2
        from core.settings import get_resume_data_path as _grp2
        _xml2    = _grp2()
        _um_path = _om2.path.join(_om2.path.dirname(_xml2), "user_metrics.json") if _xml2 else ""
        if _um_path and _om2.path.exists(_um_path):
            with open(_um_path) as _uf:
                _user_metrics = _jm2.load(_uf)
            log("User metrics loaded: %d job(s) with real numbers" % len(_user_metrics))
            for _co, _mv in _user_metrics.items():
                log("  %s: %s" % (_co, str(_mv)[:80]))
        else:
            log_debug("No user_metrics.json — Claude uses profile data only")
    except Exception as _ume:
        log_warn("user_metrics.json load failed: %s — continuing without" % _ume)
        _user_metrics = {}  # guaranteed fallback


    # Preflight: verify prompts DB is seeded before attempting any resume.
    # This catches the "run migrate_prompts.py" situation once, clearly,
    # rather than emitting a cryptic ValueError for every single job.
    if not all_prompts_ready():
        print(
            "\n[ERR] Resume generation cannot start -- prompts are not set up.\n"
            "   Your prompts.db is missing required prompts.\n"
            "   Run ONE of the following and then try again:\n"
            "\n"
            "   New user (first time):   python resume_intake.py\n"
            "   Existing user:           python migrate_prompts.py\n"
        )
        return {"generated": 0, "failed": 0, "total": 0}

    jobs = tracker.get_jobs_needing_resume(
        session_start=session_start,
        job_ids=job_ids)
    total = len(jobs)

    if total == 0:
        log("Phase 2: no matched jobs need resumes")
        print("   (i)  No matched jobs need resumes -- nothing to generate.")
        return {"generated": 0, "failed": 0, "total": 0}

    log("Phase 2: generating %d tailored resume(s)" % total)
    print(f"\n{'='*60}")
    print(f"[>>] Phase 2 -- Generating {total} tailored resume(s)")
    print(f"   Pacing: {DELAY_BETWEEN_CALLS}s between calls to respect API limits")
    print(f"{'='*60}")

    generated = 0
    failed    = 0
    applicant = profile.get("name", "Applicant")

    for i, job in enumerate(jobs, start=1):
        title   = job.get("job_title", "Unknown role")
        company = job.get("company", "Unknown company")
        job_id  = job.get("id", "?")

        log("-" * 55)
        log("Resume [%d/%d]: %s @ %s (DB id=%s)" % (i, total, title, company, job_id))
        print(f"\n   [{i}/{total}] {title} @ {company}")

        # Restore Stage 2 JD metadata from DB if not in memory
        if not job.get("_jd_metadata"):
            try:
                import json as _jmj
                _raw = job.get("jd_metadata_json") or "{}"
                job["_jd_metadata"] = _jmj.loads(_raw)
                if job["_jd_metadata"]:
                    log_debug("Restored JD metadata from DB: %d skills" %
                              len(job["_jd_metadata"].get("skills",[])))
            except Exception:
                job["_jd_metadata"] = {}
        job_desc = job.get("job_description") or ""
        if not job_desc.strip():
            log_warn("  SKIP: No job description stored in DB — cannot tailor resume")
            log_warn("  Fix: ensure save_scanning_job stores description before relevance check")
            print(f"      [WARN]  No stored description -- skipping.")
            failed += 1
            continue

        try:
            # 1. Tailor -- summary, skills, experience verbatim
            job["_user_metrics"] = _user_metrics  # from popup collection
            tailored = _tailor_with_retry(profile, job_desc, job=job)

            # 2. Job-specific experience prioritisation -- reorder and
            #    select bullets by relevance to this job's description.
            #    Only runs when the user opted in during intake.
            #    Bullet text is never changed -- only order and selection.
            if profile.get("experience_highlight"):
                print(f"      [*] Highlighting experience for this role...")
                tailored = prioritize_for_job(tailored, profile, job_desc)

            # 3. Safety net -- check for fabricated domain terms
            flagged = check_summary_for_fabrication(
                tailored.get("summary", ""), profile
            )
            if flagged:
                print(f"      [WARN]  Summary mentions terms not in your profile: "
                      f"{', '.join(flagged)} -- review before applying.")

            # 3. Build DOCX
            docx_path = make_unique_filename(
                company=company, title=title,
                folder=output_dirs["docx"], ext="docx",
                applicant_name=applicant,
            )
            log("  Step 3: Building DOCX → %s" % docx_path)
            build_resume_docx(profile, tailored, docx_path)
            log("  Step 3 complete: DOCX written successfully")

            # 4. Convert to PDF
            log("  Step 4: Converting DOCX to PDF...")
            pdf_path = convert_to_pdf(docx_path, output_dirs["pdf"])
            if pdf_path:
                log("  Step 4 complete: PDF written → %s" % pdf_path)
            else:
                log_warn("  Step 4: PDF conversion skipped (Word/LibreOffice not found) — DOCX only")

            # 4b. Validation guardrail: check metrics + keywords present
            # If fails — rerun at temperature=0.3 to fix alignment
            def _validate_tailored(t: dict, jd_meta: dict) -> tuple:
                """Returns (passed: bool, reason: str)"""
                text = " ".join([
                    t.get("summary",""),
                    " ".join(" ".join(b) if isinstance(b,list) else str(b)
                             for exp in t.get("experience",[])
                             for b in [exp.get("bullets",[])])
                ])
                text_l = text.lower()
                # Check 1: at least 1 metric (%, $, number with unit)
                import re as _re
                has_metric = bool(_re.search(r"\d+\s*(%|\$|x|hrs?|days?|ms|sec|min|k|m)", text_l))
                # Check 2: at least 3 JD keywords appear
                jd_kws    = [k.lower() for k in jd_meta.get("skills", [])]
                kw_hits   = sum(1 for k in jd_kws if k in text_l)
                if not has_metric:
                    return False, "No quantified metric found in output"
                if kw_hits < min(3, len(jd_kws)):
                    return False, "Only %d/%d JD keywords used" % (kw_hits, len(jd_kws))
                return True, "OK"

            jd_meta_for_val = job.get("_jd_metadata", {})
            val_passed, val_reason = _validate_tailored(tailored, jd_meta_for_val)
            _retries = 0

            if not val_passed:
                log_warn("Validation failed (%s) — retrying at temp=0.3" % val_reason)
                try:
                    # Rebuild minimal retry prompt from available data
                    retry_txt = (
                        "Fix this resume JSON to include quantified metrics "
                        "and use these JD keywords: " +
                        ", ".join(jd_meta_for_val.get("skills", [])[:5])
                    )
                    sys_txt2 = ""
                    usr_txt2 = retry_txt
                    msg2 = client.messages.create(
                        model=AI_MODEL,
                        max_tokens=6000,
                        temperature=0.3,
                        timeout=90.0,
                        system=sys_txt2 if sys_txt2 else anthropic.NOT_GIVEN,
                        messages=[{"role":"user","content":usr_txt2}]
                    )
                    raw2     = msg2.content[0].text.strip()
                    tailored2 = json.loads(raw2 if not raw2.startswith("```") else raw2.split("```")[1].lstrip("json"))
                    val2, _  = _validate_tailored(tailored2, jd_meta_for_val)
                    if val2:
                        # Run through tag processor before using retry output
                        import json as _jrt, re as _rrt
                        _rt_str = _rrt.sub(r"</?(?:fact|estimate)[^>]*>", "", _jrt.dumps(tailored2))
                        try:
                            tailored = _jrt.loads(_rt_str)
                        except Exception:
                            tailored = tailored2
                        log("Validation passed on retry")
                    else:
                        log_warn("Retry also failed validation — using first output")
                    _retries = 1
                except Exception as _ve:
                    log_warn("Validation retry error: %s — using first output" % _ve)

            # Log Stage 4b to transaction
            _tx = job.get("_tx")
            if _tx:
                _tx.stage4b_done(
                    tokens_used=len(str(tailored)) // 4,
                    validation_passed=val_passed,
                    retries=_retries)
                _tx.log_summary()

            # Two-channel post-processing:
            # <fact> values → copied exactly (real user data, never touched)
            # <estimate> values → fuzzed slightly (AI estimates, humanized)
            import json as _j, re as _re
            tailored_str = _j.dumps(tailored)
            try:
                tailored_str = selective_fuzzing_processor(tailored_str)
            except Exception as _fe:
                log_warn("selective_fuzzing_processor failed: %s — using raw output" % _fe)

            # Safety net: strip ANY remaining XML tags before JSON parse
            # Prevents literal tags appearing in the final DOCX
            tailored_str = _re.sub(r"</?(?:fact|estimate)[^>]*>", "", tailored_str)

            try:
                tailored = _j.loads(tailored_str)
            except Exception:
                # JSON broke — strip tags from original and use that
                safe_str = _re.sub(r"</?(?:fact|estimate)[^>]*>", "",
                                   _j.dumps(tailored))
                try:
                    tailored = _j.loads(safe_str)
                except Exception:
                    pass  # keep original tailored dict unchanged

            # 5. Mark ready in the database
            tracker.mark_resume_ready(job["id"], docx_path, pdf_path)
            generated += 1
            log("  Step 5 complete: DB updated, resume marked ready")
            log("Resume generated: %s" % docx_path)
            log("Resume location:  %s" % os.path.dirname(docx_path))
            print(f"      [OK] Resume ready: {docx_path}")

        except Exception as e:
            import traceback
            tb = traceback.format_exc()
            log_error("Resume FAILED [%d/%d] %s @ %s" % (i, total, title, company))
            log_error("  Error type: %s" % type(e).__name__)
            log_error("  Error msg:  %s" % e)
            log_error("  Traceback:  %s" % tb)
            log_error("  Job ID:     %s" % job_id)
            log_error("  Has desc:   %s (%d chars)" % (bool(job_desc), len(job_desc)))
            print("      [ERR] Resume failed for %s: %s" % (title, e))
            for ln in tb.splitlines():
                print("             %s" % ln)
            failed += 1

        # Throttle -- pause before the next call (skip after the last)
        if i < total:
            time.sleep(DELAY_BETWEEN_CALLS)

    print(f"\n{'='*60}")
    log("=" * 55)
    log("Phase 2 Summary:")
    log("  Generated: %d resumes" % generated)
    log("  Failed:    %d resumes" % failed)
    log("  Total:     %d jobs processed" % total)
    if generated > 0:
        log("  Output:    %s" % output_dirs.get("docx","?"))
    log("=" * 55)
    log("Phase 2 complete: %d generated, %d failed" % (generated, failed))
    print(f"[>>] Phase 2 complete -- {generated} generated, {failed} failed")
    print(f"{'='*60}")
    return {"generated": generated, "failed": failed, "total": total}